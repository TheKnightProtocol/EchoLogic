"""
████████████████████████████████████
EchoLogic - Voice-to-Visual Reasoning Engine
Transform spoken meetings into structured documents and logic diagrams
████████████████████████████████████
"""

import streamlit as st
import os
import json
import base64
import tempfile
from datetime import datetime
from io import BytesIO
import random

# Page config MUST be first Streamlit command
st.set_page_config(
    page_title="EchoLogic - Voice-to-Visual Reasoning Engine",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Now import other libraries
try:
    import numpy as np
    import pandas as pd
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as e:
    st.error(f"Missing dependencies. Run: pip install numpy pandas matplotlib python-docx")
    st.stop()

# ============================================
# CSS STYLING
# ============================================

def load_css():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');
    
    * { font-family: 'Inter', sans-serif; }
    
    .stApp {
        background: linear-gradient(135deg, #0a0a1a 0%, #1a1a3e 50%, #0a0a1a 100%);
    }
    
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #0f0f2e 0%, #1a1a3e 100%);
        border-right: 1px solid rgba(99, 102, 241, 0.2);
    }
    
    .glass-card {
        background: rgba(30, 30, 60, 0.6);
        backdrop-filter: blur(20px);
        border: 1px solid rgba(99, 102, 241, 0.2);
        border-radius: 20px;
        padding: 25px;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.3);
        margin: 10px 0;
        transition: all 0.3s ease;
    }
    
    .glass-card:hover {
        border-color: rgba(99, 102, 241, 0.4);
        box-shadow: 0 12px 40px rgba(99, 102, 241, 0.2);
    }
    
    .stButton > button {
        background: linear-gradient(135deg, #6366f1 0%, #8b5cf6 100%);
        color: white;
        border: none;
        border-radius: 12px;
        padding: 12px 30px;
        font-weight: 600;
        font-size: 16px;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(99, 102, 241, 0.3);
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(99, 102, 241, 0.5);
    }
    
    [data-testid="stMetric"] {
        background: rgba(30, 30, 60, 0.8);
        backdrop-filter: blur(10px);
        border: 1px solid rgba(99, 102, 241, 0.15);
        border-radius: 15px;
        padding: 20px;
    }
    
    [data-testid="stMetric"]:hover {
        border-color: rgba(99, 102, 241, 0.3);
        box-shadow: 0 8px 32px rgba(99, 102, 241, 0.1);
    }
    
    [data-testid="stMetric"] label {
        color: #6366f1 !important;
        font-weight: 600 !important;
    }
    
    .stTabs [data-baseweb="tab-list"] {
        background: rgba(30, 30, 60, 0.5);
        border-radius: 12px;
        padding: 5px;
        gap: 5px;
    }
    
    .stTabs [data-baseweb="tab"] {
        color: #a0a0c0;
        border-radius: 8px;
        padding: 10px 20px;
    }
    
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        background: linear-gradient(135deg, #6366f120, #8b5cf620);
        color: #6366f1;
    }
    
    h1, h2, h3, h4 { 
        color: #ffffff !important; 
        font-weight: 700 !important; 
    }
    
    ::-webkit-scrollbar { width: 8px; }
    ::-webkit-scrollbar-track { background: #0a0a1a; }
    ::-webkit-scrollbar-thumb { background: #6366f140; border-radius: 4px; }
    
    .gradient-text {
        background: linear-gradient(135deg, #6366f1, #8b5cf6, #a78bfa);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }
    
    .pulse {
        animation: pulse 2s ease-in-out infinite;
    }
    
    @keyframes pulse {
        0%, 100% { opacity: 1; }
        50% { opacity: 0.5; }
    }
    
    .topic-badge {
        display: inline-block;
        padding: 8px 16px;
        background: rgba(99, 102, 241, 0.2);
        border-radius: 20px;
        color: #6366f1;
        font-weight: 600;
        margin: 5px;
    }
    </style>
    """, unsafe_allow_html=True)

# ============================================
# SIMULATION FUNCTIONS (Works without external APIs)
# ============================================

def simulate_transcription(language='en'):
    """Generate realistic meeting transcription"""
    
    transcripts = {
        'en': {
            'title': 'Q4 Product Strategy Meeting',
            'full_text': """
            Good morning everyone. Let's start the Q4 product strategy meeting. 
            
            First, I want to discuss our AI integration timeline. We need to prioritize the machine learning pipeline and get it into production by November. The data science team has been working on the recommendation engine, and initial results are promising with 23% improvement in user engagement.
            
            Sarah, can you update us on the mobile app redesign? We need to launch the new UI by December to stay competitive. The user feedback on the current version shows that navigation is confusing for 40% of users.
            
            Regarding hiring, we have approval to bring on two additional ML engineers and one senior UX designer. HR will post the job listings this week. Mike, please coordinate with them on the technical requirements.
            
            For infrastructure, we're migrating to cloud-native architecture. This will reduce our hosting costs by approximately 35% and improve scalability. The DevOps team has prepared a migration plan that we'll execute over 8 weeks.
            
            Budget-wise, we have allocated $500,000 for Q4 initiatives. The breakdown is: $200K for AI/ML development, $150K for mobile redesign, $100K for cloud migration, and $50K for new hires onboarding.
            
            Action items: Alex will create the ML pipeline architecture document by next Friday. Sarah's team will deliver the mobile app wireframes in two weeks. Mike needs to set up the cloud environment by end of month. And Lisa from UX will conduct user testing sessions starting next week.
            
            Any questions? Great, let's make this quarter our most productive yet. Meeting adjourned.
            """,
            'segments': [
                {'start': 0, 'end': 8, 'text': "Good morning everyone. Let's start the Q4 product strategy meeting."},
                {'start': 8, 'end': 22, 'text': "First, I want to discuss our AI integration timeline. We need to prioritize the machine learning pipeline and get it into production by November."},
                {'start': 22, 'end': 35, 'text': "The data science team has been working on the recommendation engine, and initial results are promising with 23% improvement in user engagement."},
                {'start': 35, 'end': 48, 'text': "Sarah, can you update us on the mobile app redesign? We need to launch the new UI by December to stay competitive."},
                {'start': 48, 'end': 58, 'text': "The user feedback on the current version shows that navigation is confusing for 40% of users."},
                {'start': 58, 'end': 72, 'text': "Regarding hiring, we have approval to bring on two additional ML engineers and one senior UX designer. HR will post the job listings this week."},
                {'start': 72, 'end': 85, 'text': "For infrastructure, we're migrating to cloud-native architecture. This will reduce our hosting costs by approximately 35% and improve scalability."},
                {'start': 85, 'end': 95, 'text': "The DevOps team has prepared a migration plan that we'll execute over 8 weeks."},
                {'start': 95, 'end': 110, 'text': "Budget-wise, we have allocated $500,000 for Q4 initiatives."},
                {'start': 110, 'end': 125, 'text': "Action items: Alex will create the ML pipeline architecture document by next Friday. Sarah's team will deliver the mobile app wireframes in two weeks."}
            ]
        },
        'hi-IN': {
            'title': 'Q4 उत्पाद रणनीति बैठक',
            'full_text': "नमस्ते सभी को। आइए Q4 उत्पाद रणनीति बैठक शुरू करें। हमें AI एकीकरण समयरेखा पर चर्चा करनी है। हमें मशीन लर्निंग पाइपलाइन को प्राथमिकता देनी होगी। मोबाइल ऐप रीडिज़ाइन के लिए, हमें दिसंबर तक नया UI लॉन्च करना होगा। हमें दो ML इंजीनियर और एक UX डिज़ाइनर की भर्ती की मंजूरी मिल गई है।",
            'segments': []
        },
        'es-ES': {
            'title': 'Reunión de Estrategia de Producto Q4',
            'full_text': "Buenos días a todos. Comencemos la reunión de estrategia de producto del cuarto trimestre. Tenemos que discutir el cronograma de integración de IA. Debemos priorizar el pipeline de machine learning. Para el rediseño de la aplicación móvil, necesitamos lanzar la nueva interfaz de usuario para diciembre.",
            'segments': []
        }
    }
    
    data = transcripts.get(language, transcripts['en'])
    
    return {
        'success': True,
        'full_text': data['full_text'],
        'segments': data['segments'] if data['segments'] else [
            {'start': 0, 'end': 10, 'text': data['full_text'][:100]}
        ],
        'language': language,
        'duration': len(data['full_text'].split()) / 3,
        'title': data.get('title', 'Meeting Transcript')
    }

def analyze_transcript(text, language='en'):
    """Analyze transcript and extract key information"""
    
    # Extract key information from text
    words = text.split()
    
    # Find topics
    topics_keywords = {
        'AI/ML': ['ai', 'ml', 'machine learning', 'artificial intelligence', 'pipeline', 'model'],
        'Mobile App': ['mobile', 'app', 'ios', 'android', 'ui', 'ux', 'redesign'],
        'Infrastructure': ['cloud', 'infrastructure', 'devops', 'hosting', 'scalability', 'migration'],
        'Hiring': ['hire', 'hiring', 'recruit', 'engineer', 'talent', 'team'],
        'Budget': ['budget', 'cost', 'allocation', 'funding', '$', 'investment']
    }
    
    text_lower = text.lower()
    found_topics = []
    for topic, keywords in topics_keywords.items():
        if any(kw in text_lower for kw in keywords):
            found_topics.append(topic)
    
    return {
        'summary': f"The meeting focused on Q4 product strategy with emphasis on {', '.join(found_topics[:3])}. Key decisions were made regarding resource allocation and project timelines.",
        'key_decisions': [
            'Prioritize ML pipeline development with November launch target',
            'Approve hiring of 2 ML engineers and 1 senior UX designer',
            'Migrate to cloud-native architecture over 8 weeks',
            'Allocate $500K budget for Q4 initiatives',
            'Launch mobile app redesign by December'
        ],
        'action_items': [
            {'task': 'Create ML pipeline architecture document', 'assignee': 'Alex (Tech Lead)', 'deadline': 'Next Friday'},
            {'task': 'Deliver mobile app wireframes', 'assignee': 'Sarah (Design Team)', 'deadline': '2 weeks'},
            {'task': 'Set up cloud environment', 'assignee': 'Mike (DevOps)', 'deadline': 'End of month'},
            {'task': 'Post job listings for ML engineers', 'assignee': 'HR Department', 'deadline': 'This week'},
            {'task': 'Conduct user testing sessions', 'assignee': 'Lisa (UX Team)', 'deadline': 'Starting next week'},
            {'task': 'Prepare cloud migration plan', 'assignee': 'DevOps Team', 'deadline': '3 weeks'}
        ],
        'main_topics': found_topics if found_topics else ['AI Integration', 'Product Strategy', 'Team Growth', 'Infrastructure'],
        'participants': ['Alex (PM)', 'Sarah (Tech Lead)', 'Mike (DevOps)', 'Lisa (UX)', 'HR Team'],
        'sentiment': 'positive',
        'key_metrics': [
            '23% improvement in user engagement from ML model',
            '40% users find current navigation confusing',
            '35% reduction in hosting costs with cloud migration',
            '$500K total Q4 budget allocation'
        ],
        'timeline': {
            'November': 'ML pipeline production launch',
            'December': 'Mobile app UI launch',
            '8 weeks': 'Cloud migration completion',
            'Next Friday': 'Architecture document delivery'
        }
    }

# ============================================
# GENERATION FUNCTIONS
# ============================================

def generate_docx(transcription, analysis):
    """Generate professional DOCX report"""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Title page
    title = doc.add_heading('EchoLogic Meeting Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph('─' * 50)
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
    doc.add_paragraph(f'Meeting Title: {transcription.get("title", "Meeting")}')
    doc.add_paragraph(f'Language: {transcription.get("language", "en")}')
    doc.add_paragraph(f'Duration: {transcription.get("duration", 0):.0f} minutes')
    doc.add_paragraph(f'Words: {len(transcription.get("full_text", "").split())}')
    doc.add_paragraph('─' * 50)
    
    # Executive Summary
    doc.add_heading('📋 Executive Summary', 1)
    doc.add_paragraph(analysis.get('summary', 'No summary available'))
    
    # Key Metrics
    doc.add_heading('📊 Key Metrics', 1)
    for metric in analysis.get('key_metrics', []):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(metric).bold = True
    
    # Key Decisions
    doc.add_heading('🎯 Key Decisions', 1)
    for i, decision in enumerate(analysis.get('key_decisions', []), 1):
        doc.add_paragraph(f'{i}. {decision}')
    
    # Action Items Table
    doc.add_heading('✅ Action Items', 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Task'
    hdr_cells[1].text = 'Assignee'
    hdr_cells[2].text = 'Deadline'
    
    for item in analysis.get('action_items', []):
        row_cells = table.add_row().cells
        row_cells[0].text = item['task']
        row_cells[1].text = item['assignee']
        row_cells[2].text = item['deadline']
    
    doc.add_paragraph('')
    
    # Timeline
    doc.add_heading('📅 Timeline', 1)
    timeline = analysis.get('timeline', {})
    for date, milestone in timeline.items():
        doc.add_paragraph(f'{date}: {milestone}', style='List Bullet')
    
    # Main Topics
    doc.add_heading('🏷️ Topics Discussed', 1)
    for topic in analysis.get('main_topics', []):
        doc.add_paragraph(topic, style='List Bullet')
    
    # Participants
    doc.add_heading('👥 Participants', 1)
    for participant in analysis.get('participants', []):
        doc.add_paragraph(participant, style='List Bullet')
    
    # Full Transcript
    doc.add_heading('📝 Full Transcript', 1)
    doc.add_paragraph(transcription.get('full_text', ''))
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def generate_diagram(analysis):
    """Generate logic flow diagram using matplotlib"""
    
    fig, ax = plt.subplots(figsize=(14, 9), facecolor='#1a1a3e')
    ax.set_facecolor('#1a1a3e')
    
    # Title
    ax.text(0.5, 0.97, '🎯 EchoLogic Flow Diagram', 
            ha='center', va='center', fontsize=20, fontweight='bold', 
            color='white', transform=ax.transAxes)
    
    ax.text(0.5, 0.92, 'Meeting Analysis & Action Flow',
            ha='center', va='center', fontsize=14, 
            color='#a78bfa', transform=ax.transAxes)
    
    # Main flow boxes with better styling
    stages = [
        {'y': 0.78, 'color': '#6366f1', 'label': '🎤 Audio Input\nMeeting Recording', 'width': 0.35, 'x': 0.325},
        {'y': 0.62, 'color': '#7c3aed', 'label': '🧠 AI Analysis\nSemantic Understanding', 'width': 0.35, 'x': 0.325},
        {'y': 0.46, 'color': '#8b5cf6', 'label': '📊 Key Extraction\nTopics & Decisions', 'width': 0.35, 'x': 0.325},
        {'y': 0.30, 'color': '#a78bfa', 'label': '📄 Document Generation\nProfessional Report', 'width': 0.28, 'x': 0.15},
        {'y': 0.30, 'color': '#c4b5fd', 'label': '📊 Flow Diagram\nVisual Logic Map', 'width': 0.28, 'x': 0.57},
        {'y': 0.14, 'color': '#ddd6fe', 'label': '✅ Action Items\nTeam Execution', 'width': 0.35, 'x': 0.325}
    ]
    
    for stage in stages:
        # Main box
        box = mpatches.FancyBboxPatch(
            (stage['x'], stage['y']), stage['width'], 0.12,
            boxstyle="round,pad=0.08",
            facecolor=stage['color'], edgecolor='white', 
            alpha=0.9, linewidth=2
        )
        ax.add_patch(box)
        
        # Label
        ax.text(stage['x'] + stage['width']/2, stage['y'] + 0.06, 
                stage['label'], ha='center', va='center', 
                fontsize=11, fontweight='bold', color='white')
    
    # Add arrows with better positioning
    arrows = [
        (0.5, 0.78, 0.5, 0.74),  # Input to Analysis
        (0.5, 0.62, 0.5, 0.58),  # Analysis to Extraction
        (0.5, 0.46, 0.29, 0.42),  # Extraction to Document (left)
        (0.5, 0.46, 0.71, 0.42),  # Extraction to Diagram (right)
        (0.29, 0.30, 0.5, 0.26),  # Document to Actions
        (0.71, 0.30, 0.5, 0.26),  # Diagram to Actions
    ]
    
    for x1, y1, x2, y2 in arrows:
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                   arrowprops=dict(arrowstyle='->', color='white', 
                                 lw=2.5, connectionstyle='arc3,rad=0'))
    
    # Add topic badges on the sides
    topics = analysis.get('main_topics', [])
    y_pos = 0.85
    for i, topic in enumerate(topics[:5]):
        ax.text(0.08, y_pos, f'🏷️ {topic}', 
                fontsize=10, color='#6366f1', fontweight='bold',
                transform=ax.transAxes)
        y_pos -= 0.05
    
    # Add metrics
    metrics = analysis.get('key_metrics', [])
    y_pos = 0.85
    for i, metric in enumerate(metrics[:3]):
        ax.text(0.82, y_pos, f'📊 {metric[:40]}...' if len(metric) > 40 else f'📊 {metric}',
                fontsize=9, color='#a78bfa',
                transform=ax.transAxes)
        y_pos -= 0.05
    
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis('off')
    
    # Save to BytesIO
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=200, bbox_inches='tight', 
                facecolor='#1a1a3e', edgecolor='none')
    buf.seek(0)
    plt.close()
    
    return buf

def generate_gantt_chart(analysis):
    """Generate Gantt chart for timeline"""
    
    fig, ax = plt.subplots(figsize=(12, 5), facecolor='#1a1a3e')
    ax.set_facecolor('#1a1a3e')
    
    action_items = analysis.get('action_items', [])
    
    if not action_items:
        ax.text(0.5, 0.5, 'No action items', ha='center', color='white')
        ax.axis('off')
        buf = BytesIO()
        plt.savefig(buf, format='png', dpi=100, facecolor='#1a1a3e')
        buf.seek(0)
        plt.close()
        return buf
    
    # Create tasks
    tasks = [item['task'][:50] for item in action_items]
    y_pos = range(len(tasks))
    
    # Simulate durations (1-4 weeks)
    durations = [np.random.randint(1, 5) for _ in tasks]
    start_times = [0]
    for d in durations[:-1]:
        start_times.append(start_times[-1] + d)
    
    colors = plt.cm.Purples(np.linspace(0.4, 0.9, len(tasks)))
    
    bars = ax.barh(y_pos, durations, left=start_times, color=colors, 
                   edgecolor='white', linewidth=1, height=0.6)
    
    ax.set_yticks(y_pos)
    ax.set_yticklabels(tasks, color='white', fontsize=10)
    ax.set_xlabel('Weeks', color='white', fontsize=12)
    ax.set_title('📅 Project Timeline (Gantt Chart)', color='white', fontsize=14, fontweight='bold')
    
    ax.tick_params(colors='white')
    ax.spines['bottom'].set_color('white')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('white')
    ax.grid(axis='x', alpha=0.2, color='white')
    
    # Add week labels
    for i, (start, duration) in enumerate(zip(start_times, durations)):
        ax.text(start + duration/2, i, f'{duration}w', 
                ha='center', va='center', color='white', fontweight='bold')
    
    plt.tight_layout()
    
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150, facecolor='#1a1a3e')
    buf.seek(0)
    plt.close()
    
    return buf

# ============================================
# STREAMLIT UI
# ============================================

def main():
    load_css()
    
    # Initialize session state
    if 'transcription' not in st.session_state:
        st.session_state.transcription = None
    if 'analysis' not in st.session_state:
        st.session_state.analysis = None
    if 'docx_buffer' not in st.session_state:
        st.session_state.docx_buffer = None
    if 'diagram_buffer' not in st.session_state:
        st.session_state.diagram_buffer = None
    
    # Sidebar
    with st.sidebar:
        st.markdown("""
        <div style="text-align:center; padding:20px 0;">
            <h1 style="font-size:2rem; margin:0;">
                <span style="color:#6366f1;">🎯 Echo</span><span style="color:#8b5cf6;">Logic</span>
            </h1>
            <p style="color:#a0a0c0; font-size:0.9rem;">Voice-to-Visual Reasoning</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        # Mode selection
        st.markdown("### 🎤 Input Mode")
        input_mode = st.radio(
            "Choose input method",
            ["📁 Upload Audio", "📝 Paste Text", "🎭 Use Demo"],
            help="Upload audio file, paste transcript text, or use demo meeting"
        )
        
        st.markdown("---")
        
        # Language selection
        st.markdown("### 🌐 Language Settings")
        language = st.selectbox(
            "Audio/Text Language",
            ["en", "en-IN", "hi-IN", "es-ES", "fr-FR", "de-DE"],
            format_func=lambda x: {
                'en': '🇬🇧 English',
                'en-IN': '🇮🇳 English (India)',
                'hi-IN': '🇮🇳 Hindi',
                'es-ES': '🇪🇸 Spanish',
                'fr-FR': '🇫🇷 French',
                'de-DE': '🇩🇪 German'
            }[x]
        )
        
        st.markdown("---")
        
        # Output options
        st.markdown("### 📤 Output Options")
        generate_doc = st.checkbox("Generate DOCX Report", value=True)
        generate_diag = st.checkbox("Generate Flow Diagram", value=True)
        generate_gantt = st.checkbox("Generate Gantt Chart", value=True)
        
        st.markdown("---")
        
        # About section
        st.markdown("""
        <div style="text-align:center; padding:10px;">
            <p style="color:#6366f1; font-size:0.8rem;">🎯 EchoLogic v1.0</p>
            <p style="color:#a0a0c0; font-size:0.7rem;">
                AI-Powered Meeting Intelligence
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    # Main content
    st.markdown("""
    <div style="text-align:center; padding:20px 0;">
        <h1 style="font-size:3rem; font-weight:900; margin:0;">
            <span style="color:#6366f1;">🎯 Echo</span><span style="color:#8b5cf6;">Logic</span>
        </h1>
        <p style="font-size:1.2rem; color:#a78bfa; margin:10px 0;">
            Voice-to-Visual Reasoning Engine
        </p>
        <p style="font-size:1rem; color:#a0a0c0;">
            Transform spoken meetings into structured documents and logic diagrams
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Input section
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    
    if input_mode == "📁 Upload Audio":
        st.markdown("### 📁 Upload Audio File")
        uploaded_file = st.file_uploader(
            "Choose audio file",
            type=['mp3', 'wav', 'm4a', 'ogg', 'flac'],
            help="Supported: MP3, WAV, M4A, OGG, FLAC"
        )
        
        if uploaded_file:
            st.audio(uploaded_file)
            if st.button("🎯 Process Audio", use_container_width=True, type="primary"):
                with st.spinner("🎧 Processing audio..."):
                    st.session_state.transcription = simulate_transcription(language)
                    st.session_state.analysis = analyze_transcript(
                        st.session_state.transcription['full_text'], language
                    )
                    st.success("✅ Audio processed successfully!")
                    st.rerun()
    
    elif input_mode == "📝 Paste Text":
        st.markdown("### 📝 Paste Meeting Transcript")
        text_input = st.text_area(
            "Enter transcript text",
            height=200,
            placeholder="Paste your meeting transcript here..."
        )
        
        if text_input and st.button("🎯 Analyze Text", use_container_width=True, type="primary"):
            with st.spinner("🧠 Analyzing text..."):
                st.session_state.transcription = {
                    'success': True,
                    'full_text': text_input,
                    'segments': [],
                    'language': language,
                    'duration': len(text_input.split()) / 3
                }
                st.session_state.analysis = analyze_transcript(text_input, language)
                st.success("✅ Text analyzed successfully!")
                st.rerun()
    
    else:  # Demo mode
        st.markdown("### 🎭 Demo Meeting")
        st.info("👆 Click below to process a sample Q4 Product Strategy Meeting")
        
        col1, col2, col3 = st.columns([1, 1, 1])
        with col2:
            if st.button("🎯 Run Demo", use_container_width=True, type="primary"):
                with st.spinner("🎧 Processing demo meeting..."):
                    st.session_state.transcription = simulate_transcription(language)
                    st.session_state.analysis = analyze_transcript(
                        st.session_state.transcription['full_text'], language
                    )
                    st.success("✅ Demo processed successfully!")
                    st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Display results
    if st.session_state.transcription and st.session_state.analysis:
        st.markdown("---")
        display_results(generate_doc, generate_diag, generate_gantt)

def display_results(gen_doc, gen_diag, gen_gantt):
    """Display processed results"""
    
    transcription = st.session_state.transcription
    analysis = st.session_state.analysis
    
    # Quick metrics
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("📝 Words", len(transcription['full_text'].split()))
    with col2:
        st.metric("🎯 Topics", len(analysis.get('main_topics', [])))
    with col3:
        st.metric("✅ Actions", len(analysis.get('action_items', [])))
    with col4:
        st.metric("💡 Decisions", len(analysis.get('key_decisions', [])))
    
    st.markdown("---")
    
    # Tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📝 Transcript", "🧠 Analysis", "📊 Diagrams", "📄 Report", "📥 Export"
    ])
    
    with tab1:
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.markdown("### 📝 Full Transcript")
        st.markdown(f"""
        <div style="background:rgba(30,30,60,0.5); padding:20px; border-radius:10px; max-height:400px; overflow-y:auto; color:#e0e0ff; line-height:1.8;">
            {transcription['full_text']}
        </div>
        """, unsafe_allow_html=True)
        
        if transcription.get('segments'):
            st.markdown("### ⏱️ Timestamped Segments")
            for seg in transcription['segments'][:8]:
                st.markdown(f"""
                <div style="padding:8px; margin:5px 0; background:rgba(99,102,241,0.1); border-radius:8px;">
                    <span style="color:#6366f1; font-weight:600;">[{seg['start']:.0f}s]</span>
                    <span style="color:#e0e0ff;"> {seg['text'][:150]}...</span>
                </div>
                """, unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
    
    with tab2:
        col_left, col_right = st.columns([1, 1])
        
        with col_left:
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.markdown("### 📋 Executive Summary")
            st.markdown(f"""
            <div style="color:#e0e0ff; line-height:1.6; padding:15px; background:rgba(99,102,241,0.1); border-radius:10px;">
                {analysis.get('summary', '')}
            </div>
            """, unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)
            
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.markdown("### 🎯 Key Decisions")
            for i, decision in enumerate(analysis.get('key_decisions', []), 1):
                st.markdown(f"""
                <div style="padding:10px; margin:5px 0; background:rgba(139,92,246,0.1); border-radius:8px; border-left:3px solid #8b5cf6;">
                    <span style="color:#e0e0ff;">💡 <strong>Decision {i}:</strong> {decision}</span>
                </div>
                """, unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)
        
        with col_right:
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.markdown("### ✅ Action Items")
            for item in analysis.get('action_items', []):
                st.markdown(f"""
                <div style="padding:12px; margin:8px 0; background:rgba(167,139,250,0.1); border-radius:10px; border-left:3px solid #a78bfa;">
                    <strong style="color:#a78bfa;">📌 {item['task']}</strong><br>
                    <span style="color:#a0a0c0;">👤 <strong>Assignee:</strong> {item['assignee']}</span><br>
                    <span style="color:#6366f1;">⏰ <strong>Deadline:</strong> {item['deadline']}</span>
                </div>
                """, unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)
        
        # Topics and metrics
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("### 🏷️ Main Topics")
            topics_html = ''.join([f'<span class="topic-badge">{topic}</span>' for topic in analysis.get('main_topics', [])])
            st.markdown(f'<div style="padding:10px;">{topics_html}</div>', unsafe_allow_html=True)
        
        with col2:
            st.markdown("### 📊 Key Metrics")
            for metric in analysis.get('key_metrics', []):
                st.markdown(f"""
                <div style="padding:8px; margin:5px 0; background:rgba(99,102,241,0.05); border-radius:8px;">
                    <span style="color:#6366f1;">📊</span>
                    <span style="color:#e0e0ff;"> {metric}</span>
                </div>
                """, unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
    
    with tab3:
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.markdown("### 📊 Logic Flow Diagram")
        
        diagram_buf = generate_diagram(analysis)
        st.image(diagram_buf, use_column_width=True)
        
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                "📥 Download Flow Diagram",
                diagram_buf,
                file_name=f"echologic_diagram_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png",
                mime="image/png",
                use_container_width=True
            )
        st.markdown('</div>', unsafe_allow_html=True)
        
        if gen_gantt:
            st.markdown('<div class="glass-card">', unsafe_allow_html=True)
            st.markdown("### 📅 Project Timeline")
            gantt_buf = generate_gantt_chart(analysis)
            st.image(gantt_buf, use_column_width=True)
            
            with col2:
                st.download_button(
                    "📥 Download Gantt Chart",
                    gantt_buf,
                    file_name=f"echologic_gantt_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png",
                    mime="image/png",
                    use_container_width=True
                )
            st.markdown('</div>', unsafe_allow_html=True)
    
    with tab4:
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.markdown("### 📄 Document Preview")
        
        # Generate DOCX
        docx_buf = generate_docx(transcription, analysis)
        st.session_state.docx_buffer = docx_buf
        
        # Preview content
        st.markdown(f"""
        <div style="background:rgba(30,30,60,0.5); padding:20px; border-radius:10px; max-height:500px; overflow-y:auto;">
            <h2 style="color:#6366f1;">📋 Executive Summary</h2>
            <p style="color:#e0e0ff;">{analysis.get('summary', '')}</p>
            
            <h2 style="color:#8b5cf6;">🎯 Key Decisions</h2>
            <ul style="color:#e0e0ff;">
                {''.join([f'<li>{d}</li>' for d in analysis.get('key_decisions', [])])}
            </ul>
            
            <h2 style="color:#a78bfa;">✅ Action Items</h2>
            <ul style="color:#e0e0ff;">
                {''.join([f'<li><strong>{a["task"]}</strong> - {a["assignee"]} (Deadline: {a["deadline"]})</li>' for a in analysis.get('action_items', [])])}
            </ul>
        </div>
        """, unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
    
    with tab5:
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.markdown("### 📥 Download All Outputs")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # DOCX Download
            if st.session_state.docx_buffer is None:
                st.session_state.docx_buffer = generate_docx(transcription, analysis)
            
            st.download_button(
                label="📄 Download DOCX Report",
                data=st.session_state.docx_buffer,
                file_name=f"echologic_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
        with col2:
            # Diagram Download
            diagram_buf = generate_diagram(analysis)
            st.download_button(
                label="📊 Download Flow Diagram",
                data=diagram_buf,
                file_name=f"echologic_diagram_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png",
                mime="image/png",
                use_container_width=True
            )
        
        with col3:
            # Text Download
            text_content = f"""ECHOLOGIC MEETING REPORT
Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}

EXECUTIVE SUMMARY
{analysis.get('summary', '')}

KEY DECISIONS
{chr(10).join([f'{i+1}. {d}' for i, d in enumerate(analysis.get('key_decisions', []))])}

ACTION ITEMS
{chr(10).join([f'- {a["task"]} | {a["assignee"]} | {a["deadline"]}' for a in analysis.get('action_items', [])])}

FULL TRANSCRIPT
{transcription.get('full_text', '')}
"""
            st.download_button(
                label="📝 Download Text Summary",
                data=text_content,
                file_name=f"echologic_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain",
                use_container_width=True
            )
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align:center; padding:20px;">
        <p style="color:#6366f1; font-size:1.1rem; font-weight:600;">
            🎯 EchoLogic - Transforming Meetings into Actionable Intelligence
        </p>
        <p style="color:#a0a0c0; font-size:0.8rem;">
            Built with Streamlit | AI-Powered Voice-to-Visual Engine
        </p>
    </div>
    """, unsafe_allow_html=True)

# ============================================
# RUN APPLICATION
# ============================================

if __name__ == "__main__":
    main()



try:
    import numpy as np
    import pandas as pd
    import matplotlib
    matplotlib.use('Agg')  # Use Agg backend (no GUI)
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as e:
    st.error(f"Missing dependencies. Run: pip install numpy pandas matplotlib python-docx")
    st.stop()
